import csv
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import defaultdict
import matplotlib.pyplot as plt
import os
import argparse

# ---------------- Setup CLI ----------------
parser = argparse.ArgumentParser(description='Generate vulnerability report from CSV.')
parser.add_argument('csv_file', type=str, help='Path to the CSV file containing vulnerability data')
args = parser.parse_args()

# Output DOCX path
output_file = args.csv_file.replace('.csv', '.docx')

# ---------------- Word doc ----------------
doc = Document()

# ----- Font helper (keep Aptos) -----
def set_font(run, size=11, bold=False, color=None):
    run.font.name = 'Aptos'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')  # Ensure compatibility with different locales
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def add_page_break():
    doc.add_page_break()

# ----- Risk color helper (unchanged) -----
def get_impact_color(impact):
    colors = {
        'Critical': RGBColor(139, 0, 0),  # Burgundy 
        'High': RGBColor(255, 0, 0),      # Red 
        'Medium': RGBColor(255, 165, 0),  # Orange 
        'Low': RGBColor(0, 100, 0)        # Dark Green 
    }
    return colors.get(impact, RGBColor(0, 0, 0))  # Default is black

# Optional helper to safely pull a field with fallbacks
def get_field(row, *candidates):
    for key in candidates:
        if key in row and row[key] and str(row[key]).strip().lower() != 'n/a':
            return row[key]
    return ''

# ---------------- Scope ----------------
scope_heading = doc.add_paragraph('Scope', style='Heading 2')
scope_text = "The scope of the exercise includes and is limited to the following IP addresses:"
scope_paragraph = doc.add_paragraph(scope_text)
set_font(scope_paragraph.runs[0], size=11)

# Collect unique Hosts (instead of IP Address)
unique_hosts = set()
with open(args.csv_file, newline='', encoding='utf-8') as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        if 'Host' in row and row['Host']:
            unique_hosts.add(row['Host'])

unique_hosts = sorted(unique_hosts)

# 4-column table for hosts
num_columns = 4
num_rows = len(unique_hosts) // num_columns + (1 if len(unique_hosts) % num_columns != 0 else 0)
ip_table = doc.add_table(rows=max(num_rows, 1), cols=num_columns)
ip_table.style = 'Table Grid'
ip_table.alignment = WD_TABLE_ALIGNMENT.CENTER

row_index = 0
col_index = 0
for ip in unique_hosts:
    ip_table.cell(row_index, col_index).text = ip
    if ip_table.cell(row_index, col_index).paragraphs[0].runs:
        set_font(ip_table.cell(row_index, col_index).paragraphs[0].runs[0])
    col_index += 1
    if col_index == num_columns:
        col_index = 0
        row_index += 1

# ---------------- Grouping (use Risk/Host/Name) ----------------
vulnerabilities = defaultdict(list)
impact_count = defaultdict(int)
impact_unique_set = defaultdict(set)

with open(args.csv_file, newline='', encoding='utf-8') as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        risk = row.get('Risk', '').strip()
        name = row.get('Name', '').strip()
        host = row.get('Host', '').strip()
        port = row.get('Port', '').strip()

        if risk and risk.lower() != 'none' and name:
            vulnerabilities[name].append((host, port, risk, row))
            impact_count[risk] += 1
            impact_unique_set[risk].add(name)

# Sort by Risk order
impact_order = ['Critical', 'High', 'Medium', 'Low']
def risk_index(v):
    try:
        return impact_order.index(v)
    except ValueError:
        return len(impact_order)

sorted_vulnerabilities = sorted(
    vulnerabilities.items(),
    key=lambda kv: risk_index(kv[1][0][2])
)

# ---------------- Pie chart (unique Name per Risk) ----------------
df = pd.read_csv(args.csv_file)

# Use Risk/Name here (instead of Severity/Plugin Name)
impact_unique_count = {
    'Critical': df[df.get('Risk', pd.Series([])) == 'Critical'].get('Name', pd.Series([])).nunique(),
    'High':     df[df.get('Risk', pd.Series([])) == 'High'].get('Name', pd.Series([])).nunique(),
    'Medium':   df[df.get('Risk', pd.Series([])) == 'Medium'].get('Name', pd.Series([])).nunique(),
    'Low':      df[df.get('Risk', pd.Series([])) == 'Low'].get('Name', pd.Series([])).nunique()
}

labels = ['Critical', 'High', 'Medium', 'Low']
sizes = [impact_unique_count[label] for label in labels]
colors = ['#8B0000', '#FF0000', '#FFA500', '#006400']

plt.pie(sizes, labels=None, colors=colors, startangle=90,
        autopct=lambda p: f'{int(p * sum(sizes) / 100)}')
plt.title('Vulnerability Assessment')
plt.legend(labels, loc="best")
plt.axis('equal')

image_path = "impact_pie_chart.png"
plt.savefig(image_path)
plt.close()

# Insert chart + page break
doc.add_picture(image_path, width=Inches(5))
add_page_break()

# Save/cleanup (same place as قبل الجزء الملخّص)
doc.save(output_file)
os.remove(image_path)

# ---------------- Summary table ----------------
heading = doc.add_paragraph('Table of All Vulnerabilities', style='Heading 1')
set_font(heading.runs[0], size=18, bold=True)

summary_table = doc.add_table(rows=1, cols=3)
summary_table.style = 'Table Grid'

hdr_cells = summary_table.rows[0].cells
hdr_cells[0].text = '#'
hdr_cells[1].text = 'Finding'
hdr_cells[2].text = 'Risk'

for i, cell in enumerate(hdr_cells):
    if cell.paragraphs[0].runs:
        set_font(cell.paragraphs[0].runs[0], bold=True, color=RGBColor(255, 255, 255))
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '000000')
    cell._element.get_or_add_tcPr().append(shading_elm)
    if i == 2:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

vuln_number = 1
for vuln_name, hosts_ports_data in sorted_vulnerabilities:
    row_cells = summary_table.add_row().cells
    row_cells[0].text = str(vuln_number)
    if row_cells[0].paragraphs[0].runs:
        set_font(row_cells[0].paragraphs[0].runs[0])

    row_cells[1].text = vuln_name
    if row_cells[1].paragraphs[0].runs:
        set_font(row_cells[1].paragraphs[0].runs[0])

    impact_level = hosts_ports_data[0][2]  # Risk
    row_cells[2].text = impact_level
    if row_cells[2].paragraphs[0].runs:
        impact_run = row_cells[2].paragraphs[0].runs[0]
        set_font(impact_run, bold=True, color=RGBColor(255, 255, 255))
    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), {
        'Critical': '8B0000',
        'High': 'FF0000',
        'Medium': 'FFA500',
        'Low': '006400'
    }.get(impact_level, '000000'))
    row_cells[2]._element.get_or_add_tcPr().append(shading_elm)

    vuln_number += 1

add_page_break()

# ---------------- Details per finding ----------------
vuln_number = 1
for vuln_name, hosts_ports_data in sorted_vulnerabilities:
    first_row = hosts_ports_data[0][3]
    affected_hosts = {(host, port) for host, port, _, _ in hosts_ports_data}

    # Title uses Name (dynamic)
    vuln_title = doc.add_paragraph(f"{vuln_number}. {first_row.get('Name','')}", style='Heading 2')
    if vuln_title.runs:
        set_font(vuln_title.runs[0], size=14, bold=True, color=RGBColor(0, 0, 255))
    vuln_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(4.5)

    # helper to add KV row with optional risk color
    def add_row_if_data(label, data, impact=None):
        if data and str(data).strip():
            row_cells = table.add_row().cells
            row_cells[0].text = label
            if row_cells[0].paragraphs[0].runs:
                set_font(row_cells[0].paragraphs[0].runs[0], bold=True)

            data_run = row_cells[1].paragraphs[0].add_run(str(data))
            if impact:
                data_run.font.color.rgb = get_impact_color(impact)
            set_font(data_run)
            set_text_direction_ltr(row_cells[1])

    # ensure LTR for the value cell
    def set_text_direction_ltr(cell):
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '0')
        tcPr.append(bidi)

    impact = first_row.get('Risk', '')
    add_row_if_data("Risk", impact if impact else 'n/a', impact)

    affected_hosts = list(set(affected_hosts))
    affected_hosts_text = ', '.join([f"{host}:{port}" if port else f"{host}" for host, port in affected_hosts])
    add_row_if_data("Affected Hosts", affected_hosts_text if affected_hosts_text else 'n/a')

    # References / CVE if available
    add_row_if_data("References", get_field(first_row, 'CVE'))

    # Description (from Description or Synopsis)
    desc_text = get_field(first_row, 'Description', 'Synopsis')
    if desc_text:
        description_row = table.add_row().cells
        description_row[0].merge(description_row[1])
        description_row[0].text = "Description"
        if description_row[0].paragraphs[0].runs:
            set_font(description_row[0].paragraphs[0].runs[0], bold=True)
        description_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        desc_content_row = table.add_row().cells
        desc_content_row[0].merge(desc_content_row[1])
        desc_content_row[0].text = desc_text
        if desc_content_row[0].paragraphs[0].runs:
            set_font(desc_content_row[0].paragraphs[0].runs[0])
        desc_content_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        desc_content_row[0].paragraphs[0].paragraph_format.space_before = Pt(10)
        desc_content_row[0].paragraphs[0].paragraph_format.space_after = Pt(10)

    # Recommendations (Steps to Remediate OR Solution)
    fix_text = get_field(first_row, 'Steps to Remediate', 'Solution')
    if fix_text:
        recommendation_row = table.add_row().cells
        recommendation_row[0].merge(recommendation_row[1])
        recommendation_row[0].text = "Recommendations"
        if recommendation_row[0].paragraphs[0].runs:
            set_font(recommendation_row[0].paragraphs[0].runs[0], bold=True)
        recommendation_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        recommendation_content_row = table.add_row().cells
        recommendation_content_row[0].merge(recommendation_content_row[1])
        recommendation_content_row[0].text = fix_text
        if recommendation_content_row[0].paragraphs[0].runs:
            set_font(recommendation_content_row[0].paragraphs[0].runs[0])
        recommendation_content_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        recommendation_content_row[0].paragraphs[0].paragraph_format.space_before = Pt(10)
        recommendation_content_row[0].paragraphs[0].paragraph_format.space_after = Pt(10)

    # Keep same cell-coloring logic for risk tokens
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            if cell.text in ['Critical', 'High', 'Medium', 'Low']:
                cell_shading = OxmlElement('w:shd')
                cell_shading.set(qn('w:fill'), {
                    'Critical': '8B0000',
                    'High': 'FF0000',
                    'Medium': 'FFA500',
                    'Low': '006400'
                }[cell.text])
                cell._element.get_or_add_tcPr().append(cell_shading)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_font(run, color=RGBColor(255, 255, 255))
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_page_break()
    vuln_number += 1

# TOC notice + save
doc.add_paragraph('To update the Table of Contents, select it and press F9.', style='Normal')
doc.save(output_file)

print(f"Summary table, pie chart, and vulnerabilities added to {output_file}.")
